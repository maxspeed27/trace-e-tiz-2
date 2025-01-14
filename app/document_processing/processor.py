from typing import List, Optional, Dict, Any, Tuple
import fitz  # PyMuPDF
import uuid
import logging
from datetime import datetime
import re
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from ..retrieval.storage import VectorStorage
from ..models.document import DocumentMetadata, TextChunk, HierarchyContext  # Import both models
import cohere
import os
import tiktoken
from difflib import SequenceMatcher

logger = logging.getLogger(__name__)

class DocumentProcessor:
    _instance = None
    _document_hierarchy = {}  # Make this a class variable
    
    def __new__(cls, *args, **kwargs):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        if not hasattr(self, 'initialized'):
            self.chunk_size = chunk_size
            self.chunk_overlap = chunk_overlap
            
            # Use single VectorStorage instance
            self.retrieval_service = VectorStorage()
            
            # Use Cohere for embeddings
            self.cohere_client = cohere.Client(os.getenv("COHERE_API_KEY"))
            
            # Text splitter for chunking
            self.encoding = tiktoken.get_encoding("cl100k_base")
            self.text_splitter = self._split_text
            
            self.initialized = True

    async def process_document(self, file_path: str, metadata: DocumentMetadata) -> List[TextChunk]:
        logger.debug("\n=== PROCESSING DOCUMENT ===")
        logger.debug(f"Processing file: {file_path}")
        logger.debug(f"Metadata: {metadata}")
        
        doc = None
        try:
            logger.debug("Opening document with PyMuPDF...")
            doc = fitz.open(file_path)
            logger.debug(f"Successfully opened document. Pages: {len(doc)}")
            
            # Extract text and get document type hints
            text, doc_type_hints = self._extract_text(doc)
            logger.debug(f"Text extraction complete. Length: {len(text)}")
            logger.debug(f"Document type hints found: {doc_type_hints}")
            
            # Create chunks with proper metadata
            logger.debug("Starting text chunking...")
            chunks = []
            texts = self.text_splitter(text)
            logger.debug(f"Created {len(texts)} chunks")
            
            # Batch process chunks to avoid memory issues
            BATCH_SIZE = 10
            for i in range(0, len(texts), BATCH_SIZE):
                batch = texts[i:i + BATCH_SIZE]
                for j, chunk_text in enumerate(batch):
                    chunk_num = i + j + 1
                    logger.debug(f"Processing chunk {chunk_num}/{len(texts)}")
                    
                    chunk = TextChunk(
                        text=chunk_text,
                        metadata=metadata,
                        page_number=1,  # Default to 1, improve this if needed
                        chunk_id=str(uuid.uuid4()),
                        section_id=str(uuid.uuid4()),
                        level_in_document=metadata.level,
                        parent_id=None
                    )
                    chunks.append(chunk)
                    
            logger.debug(f"Successfully processed {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}", exc_info=True)
            raise
        finally:
            if doc:
                doc.close()

    async def process_document_late_chunking(self, file_path: str, metadata: DocumentMetadata) -> tuple[List[TextChunk], List[TextChunk]]:
        """Process document with hierarchical and semantic chunking"""
        print("\n=== PROCESSING DOCUMENT ===")
        print(f"Processing file: {file_path}")
        print(f"Metadata: {metadata}")
        
        doc = None
        try:
            # Extract and process text first without storing chunks
            if file_path.lower().endswith('.docx'):
                text = self._process_docx(file_path)
                document_type_hints = []
            else:
                doc = fitz.open(file_path)
                if len(doc) == 0:
                    raise ValueError(f"PDF has no pages: {file_path}")
                
                if self._needs_ocr(doc):
                    text = await self._perform_ocr(file_path)
                    document_type_hints = []
                else:
                    text, document_type_hints = self._extract_text_with_structure(doc)

            print(f"\nDocument type hints found: {document_type_hints}")
            
            # Detect document type and relationships
            doc_type, level = self._detect_document_type(text, metadata.name, document_type_hints)
            print(f"Detected document type: {doc_type}, level: {level}")

            # Update metadata with detected type and level
            metadata.document_type = doc_type
            metadata.level = level

            # Create section-level chunks
            sections = []
            current_section = {"text": "", "start_page": 1}
            
            # Split into sections based on page markers
            for line in text.split('\n'):
                if line.startswith('=== Page '):
                    if current_section["text"].strip():
                        sections.append(current_section)
                    page_num = int(line.replace('=== Page ', '').replace(' ===', ''))
                    current_section = {"text": "", "start_page": page_num}
                else:
                    current_section["text"] += line + "\n"
            
            # Add the last section
            if current_section["text"].strip():
                sections.append(current_section)

            # Create section-level chunks
            section_chunks = [
                TextChunk(
                    text=section["text"],
                    metadata=metadata,
                    page_number=section["start_page"],
                    chunk_id=str(uuid.uuid4()),
                    section_id=str(uuid.uuid4()),
                    section_reference=f"Section {i+1}",
                    level_in_document=metadata.level,
                    parent_id=None
                ) for i, section in enumerate(sections)
            ]
            
            # Create detail chunks using the text splitter
            detail_chunks = []
            for section_chunk in section_chunks:
                sub_texts = self.text_splitter(section_chunk.text)
                
                sub_chunks = [
                    TextChunk(
                        text=chunk_text,
                        metadata=metadata,
                        page_number=section_chunk.page_number,
                        chunk_id=str(uuid.uuid4()),
                        section_id=section_chunk.section_id,
                        section_reference=section_chunk.section_reference,
                        level_in_document=metadata.level,
                        parent_id=None
                    ) for chunk_text in sub_texts
                ]
                detail_chunks.extend(sub_chunks)

            print(f"Created {len(section_chunks)} section chunks and {len(detail_chunks)} detail chunks")
            return section_chunks, detail_chunks

        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}", exc_info=True)
            raise

        finally:
            if doc and hasattr(doc, 'close'):
                doc.close()

    def _create_sub_chunks(self, text: str, metadata: DocumentMetadata, page_number: int) -> List[TextChunk]:
        """Create smaller chunks from a section of text"""
        # Use the sentence splitter to create sub-chunks
        texts = self.text_splitter.split_text(text)
        
        return [
            TextChunk(
                text=chunk_text,
                metadata=metadata,
                page_number=page_number,
                chunk_id=str(uuid.uuid4()),
                level_in_document=0,
                section_id=None,
                section_reference=None,
                parent_id=None
            ) for chunk_text in texts
        ]

    def _detect_document_type(self, text: str, filename: str, document_type_hints: List[str] = None) -> tuple[str, int]:
        """Attempt to automatically detect document type and level through comprehensive analysis"""
        text_lower = text[:150].lower()  # Only look at first 150 chars
        filename_lower = filename.lower()
        
        print("\n=== Document Type Detection ===")
        print(f"Analyzing filename: {filename}")
        print(f"Document type hints: {document_type_hints}")
        print(f"First 150 chars of text: {text_lower}")
        
        # Score-based classification
        type_scores = {
            "master": 0,
            "amendment": 0,
            "sow": 0,
            "change_order": 0
        }
        
        # Check filename first - this should have highest weight
        if any(term in filename_lower for term in ["master", "msa", "_mst_"]):
            type_scores["master"] += 19
            print("Added 19 points to master from filename")
        elif any(term in filename_lower for term in ["amendment", "amd", "addendum"]):
            type_scores["amendment"] += 20
            print("Added 20 points to amendment from filename")
        elif any(term in filename_lower for term in ["sow", "statement"]):
            type_scores["sow"] += 20
            print("Added 20 points to sow from filename")
        elif any(term in filename_lower for term in ["change", "co_"]):
            type_scores["change_order"] += 20
            print("Added 20 points to change_order from filename")
        
        print(f"\nScores after filename check: {type_scores}")
        
        # Add scores from document type hints (these come from first few blocks of first page)
        if document_type_hints:
            for hint in document_type_hints:
                hint_lower = hint.lower()
                print(f"Processing hint: {hint_lower}")
                
                # Check if this is a reference to a master agreement
                is_reference_to_master = any(ref in hint_lower for ref in [
                    "to master", "to the master", "pursuant to master", 
                    "under master", "reference to master"
                ])
                
                if "amendment" in hint_lower or "addendum" in hint_lower:
                    type_scores["amendment"] += 10
                    print("Added 10 points to amendment from hint")
                
                # Handle master references after amendment check
                if "master" in hint_lower:
                    if is_reference_to_master:
                        type_scores["amendment"] += 5
                        print("Added 5 bonus points to amendment for master reference")
                    else:
                        type_scores["master"] += 8
                        print("Added 8 points to master from hint")
                elif "sow" in hint_lower or "statement of work" in hint_lower:
                    type_scores["sow"] += 10
                    print("Added 10 points to sow from hint")
                elif "change order" in hint_lower:
                    type_scores["change_order"] += 10
                    print("Added 10 points to change_order from hint")
        
        print(f"\nFinal scores: {type_scores}")
        
        # Get the type with highest score
        max_score = max(type_scores.values())
        if max_score == 0:
            # Check filename for hints if no type detected from content
            lower_filename = filename.lower()
            if any(x in lower_filename for x in ['sow', 'statement', 'work']):
                doc_type = "sow"
            elif any(x in lower_filename for x in ['amend', 'amendment']):
                doc_type = "amendment"  
            elif any(x in lower_filename for x in ['change', 'order', 'co']):
                doc_type = "change_order"
            else:
                print("No document type detected, defaulting to master agreement")
                doc_type = "master"
        else:
            doc_type = max(type_scores.items(), key=lambda x: x[1])[0]
        
        # Map to levels
        level_map = {
            "master": 1,
            "amendment": 2,
            "sow": 3,
            "change_order": 4,
            "unknown": 1  # Default unknown docs to level 1
        }
        
        level = level_map.get(doc_type, 1)  # Default to level 1 if type not in map
        print(f"\nFinal detection result - Type: {doc_type}, Level: {level}")
        
        return doc_type, level

    async def detect_parent_document(self, text: str, contract_set_id: str) -> Optional[str]:
        """Detect parent document reference from text"""
        # Common patterns for referencing master agreements
        reference_patterns = [
            r"reference.*to.*master.*agreement.*dated\s+(\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})",
            r"pursuant.*to.*master.*agreement.*dated\s+(\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})",
            r"amendment.*to.*(MSA|Master\s+Services?\s+Agreement).*dated\s+(\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})"
        ]
        
        # Search for references
        for pattern in reference_patterns:
            if match := re.search(pattern, text.lower()):
                # Query using VectorStorage instead of direct Qdrant access
                referenced_date = match.group(1)
                results = await self.retrieval_service.search_by_metadata({
                    "contract_set_id": contract_set_id,
                    "document_type": "master",
                    "effective_date": referenced_date
                }, limit=1)
                if results:
                    return results[0]["metadata"]["document_id"]
        return None 

    def _extract_version_info(self, text: str) -> tuple[Optional[str], Optional[str]]:
        """Extract version and effective date from document text"""
        version_patterns = [
            r"Version\s*[:.]?\s*(\d+(\.\d+)*)",
            r"Revision\s*[:.]?\s*(\d+(\.\d+)*)",
            r"v(\d+(\.\d+)*)"
        ]
        
        date_patterns = [
            r"Effective\s+Date\s*[:.]?\s*(\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})",
            r"Last\s+Updated\s*[:.]?\s*(\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})",
            r"Date\s*[:.]?\s*(\d{1,2}[-/.]\d{1,2}[-/.]\d{2,4})"
        ]
        
        version = next((match.group(1) for pattern in version_patterns 
                       if (match := re.search(pattern, text))), None)
        
        effective_date = next((match.group(1) for pattern in date_patterns 
                             if (match := re.search(pattern, text))), None)
        
        return version, effective_date 

    def _needs_ocr(self, doc: fitz.Document) -> bool:
        """Check if document needs OCR by looking for meaningful text content"""
        sample_text = ""
        # Check first 3 pages or all pages if less
        for page_num in range(min(3, len(doc))):
            sample_text += doc[page_num].get_text()
        
        # If we find very little text or mostly special characters, probably needs OCR
        meaningful_text = re.sub(r'[\W\d_]+', '', sample_text)
        return len(meaningful_text) < 100  # Adjust threshold as needed

    async def _perform_ocr(self, file_path: str) -> str:
        """Convert PDF to images and perform OCR"""
        images = convert_from_path(file_path)
        text = ""
        
        for i, image in enumerate(images):
            # Perform OCR on each page
            page_text = pytesseract.image_to_string(image)
            text += f"\n=== Page {i+1} ===\n{page_text}"
            
        return text 

    def _process_docx(self, file_path: str) -> str:
        """Extract text from DOCX while preserving structure"""
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            # Check if it's a heading
            if paragraph.style.name.startswith('Heading'):
                text += f"# {paragraph.text}\n"
            else:
                text += f"{paragraph.text}\n"
                
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                text += " | ".join(cell.text for cell in row.cells)
                text += "\n"
                
        return text 

    def _extract_text_with_structure(self, doc: fitz.Document) -> tuple[str, List[str]]:
        """Extract text from PDF while preserving structure and return document type hints"""
        text = ""
        document_type_hints = []
        
        # Only look at first page
        if len(doc) > 0:
            page = doc[0]
            blocks = page.get_text("dict")["blocks"]
            blocks.sort(key=lambda b: (-b["bbox"][1], b["bbox"][0]))
            
            # Only get first 2-3 blocks from top of first page
            for block in blocks[:3]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_text = " ".join(span["text"] for span in line["spans"])
                        # Look for key terms in headers
                        if any(term in line_text.lower() for term in ["master", "msa", "amendment", "addendum", "sow", "statement of work", "change order"]):
                            document_type_hints.append(line_text)
                            
        # Continue with regular text extraction for the whole document
        for page_num, page in enumerate(doc):
            text += f"\n=== Page {page_num + 1} ===\n"
            blocks = page.get_text("dict")["blocks"]
            blocks.sort(key=lambda b: (-b["bbox"][1], b["bbox"][0]))
            
            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text += span["text"] + " "
                        text += "\n"
                    text += "\n"
        
        return text.strip(), document_type_hints

    def _get_page_number(self, text: str, doc: fitz.Document) -> int:
        """Find the page number for a chunk of text using fuzzy matching"""
        text = text.strip()
        best_ratio = 0
        best_page = 1
        
        # Clean the text for comparison
        clean_text = ' '.join(text.split())
        
        for page_num in range(len(doc)):
            page_text = doc[page_num].get_text()
            clean_page_text = ' '.join(page_text.split())
            
            # Use sliding window to find best match
            window_size = len(clean_text)
            max_ratio = 0
            
            for i in range(len(clean_page_text) - window_size + 1):
                window = clean_page_text[i:i + window_size]
                ratio = SequenceMatcher(None, clean_text, window).ratio()
                max_ratio = max(max_ratio, ratio)
                
            if max_ratio > best_ratio:
                best_ratio = max_ratio
                best_page = page_num + 1
        
        logger.debug(f"Best match ratio: {best_ratio} for page {best_page}")
        
        # Only use the match if it's reasonably good
        if best_ratio > 0.8:
            return best_page
        
        logger.warning(f"No good match found (best ratio: {best_ratio}), defaulting to page 1")
        return 1

    def _split_text(self, text: str) -> List[str]:
        """Split text into chunks using tiktoken"""
        tokens = self.encoding.encode(text)
        chunks = []
        current_chunk = []
        current_size = 0
        
        for token in tokens:
            if current_size >= self.chunk_size:
                chunks.append(self.encoding.decode(current_chunk))
                current_chunk = current_chunk[-self.chunk_overlap:]
                current_size = len(current_chunk)
            current_chunk.append(token)
            current_size += 1
            
        if current_chunk:
            chunks.append(self.encoding.decode(current_chunk))
        return chunks

    def _extract_section_reference(self, text: str) -> Optional[str]:
        """Extract section reference from chunk text"""
        section_patterns = [
            r"Section\s+(\d+(\.\d+)*)",
            r"Article\s+(\d+(\.\d+)*)",
            r"Clause\s+(\d+(\.\d+)*)"
        ]
        
        for pattern in section_patterns:
            if match := re.search(pattern, text):
                return match.group(0)
        return None

    def _extract_text(self, doc: fitz.Document) -> Tuple[str, List[str]]:
        """Extract text and document type hints from PDF"""
        logger.debug("Starting text extraction...")
        
        full_text = []
        doc_type_hints = []
        
        try:
            for page_num in range(len(doc)):
                logger.debug(f"Processing page {page_num + 1}/{len(doc)}")
                page = doc[page_num]
                
                # Extract text from page
                page_text = page.get_text()
                full_text.append(f"=== page {page_num + 1} ===\n{page_text}")
                
                # Look for document type hints
                type_indicators = [
                    "Master Agreement",
                    "Amendment",
                    "Statement of Work",
                    "SOW",
                    "Change Order",
                    "Exhibit",
                    "Appendix"
                ]
                
                for indicator in type_indicators:
                    if indicator.lower() in page_text.lower():
                        doc_type_hints.append(indicator)
            
            combined_text = "\n".join(full_text)
            logger.debug(f"Text extraction complete. Total length: {len(combined_text)}")
            logger.debug(f"Document type hints found: {doc_type_hints}")
            
            return combined_text, list(set(doc_type_hints))  # Deduplicate hints
            
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}", exc_info=True)
            raise